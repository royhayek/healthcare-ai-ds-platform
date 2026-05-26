"""Text extraction from case brief files (PDF, DOCX, TXT, MD).

Supports:
  - PDF  → pdfplumber (preserves text layout, handles multi-column)
  - DOCX → python-docx (paragraphs + tables)
  - TXT / MD / any text file → UTF-8 decode with fallback to latin-1

All functions are synchronous - callers run them in a thread pool via
asyncio.to_thread() since pdfplumber and python-docx are blocking.
"""

from __future__ import annotations

import io
import logging

logger = logging.getLogger(__name__)

# File extensions we accept as case brief uploads
ACCEPTED_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
    "text/markdown",
    "text/x-markdown",
}

ACCEPTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".text"}


def extract_text_from_bytes(filename: str, content: bytes) -> str:
    """Dispatch to the right extractor based on file extension."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return _extract_pdf(content)
    if lower.endswith(".docx"):
        return _extract_docx(content)
    # TXT / MD / any other text file
    return _extract_text(content)


def _extract_pdf(content: bytes) -> str:
    try:
        import pdfplumber

        pages: list[str] = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text.strip())
        return "\n\n".join(pages)
    except Exception as exc:
        logger.warning("pdfplumber extraction failed: %s", exc)
        return ""


def _extract_docx(content: bytes) -> str:
    try:
        from docx import Document  # python-docx

        doc = Document(io.BytesIO(content))
        paragraphs: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        # Include table cell text
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    paragraphs.append(" | ".join(row_cells))
        return "\n\n".join(paragraphs)
    except Exception as exc:
        logger.warning("python-docx extraction failed: %s", exc)
        return ""


def _extract_text(content: bytes) -> str:
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("latin-1", errors="replace")
